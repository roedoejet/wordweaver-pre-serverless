import json
from tempfile import mkstemp
import requests
import csv

from docx import Document
from typing import List
from pydantic import BaseModel
from enum import Enum

from wwapi.data import URL
from wwapi.models import Response, Tier


class FileSettings(BaseModel):
    heading: str = 'Conjugations'
    headers: bool = True


class File:
    """[summary]
    """

    def __init__(self, conjugations: Response = None, tiers: List[Tier] = None):
        self._conjugations = conjugations
        self._tiers = tiers
        self.format()

    @property
    def conjugations(self):
        """
        """
        return self._conjugations

    @property
    def tiers(self):
        """
        """
        return self._tiers

    @property
    def formatted_data(self):
        """[summary]
        """
        return self._formatted_data

    def format(self):
        """[summary]

        Returns:
            [type]: [description]
        """
        self._formatted_data = []
        for conjugation in self.conjugations:
            tiered_conjugation = []
            for tier in self.tiers:
                tier = tier.dict()
                # Filter empty and sort by position
                output = sorted([x for x in conjugation['output']
                                if x[tier['key']]], key=lambda x: x['position'])
                # Join with separator
                output = tier['separator'].join(
                    map(lambda x: x[tier['key']], output))
                tiered_conjugation.append({'name': tier['name'],
                                           'options': tier['options'],
                                           'output': output})
            self._formatted_data.append(tiered_conjugation)


class DocxFile(File):
    def __init__(self, conjugations, tiers, settings: FileSettings):
        self.document = Document()
        super(DocxFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()

    def write_to_temp(self):
        # Add header
        self.document.add_heading(self.settings['heading'], 0)
        # Add each tier
        for conjugation in self.formatted_data:
            tiers = '\n'.join(map(lambda x: x['output'], conjugation))
            self.document.add_paragraph(tiers, style="List Number")
            self.document.add_paragraph()
        fd, path = mkstemp()
        self.document.save(path)
        return path


class LatexFile(File):
    def __init__(self, conjugations, tiers, settings):
        super(LatexFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()

    def write_to_temp(self):
        pass


class CsvFile(File):
    def __init__(self, conjugations, tiers, settings: FileSettings):
        super(CsvFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()

    def write_to_temp(self):
        fd, path = mkstemp()
        with open(path, 'w') as f:
            writer = csv.writer(f)
            if self.settings['headers']:
                writer.writerow([x.dict()['name'] for x in self.tiers])
            for conjugation in self.formatted_data:
                writer.writerow([x['output'] for x in conjugation])
        return path


def find(db_name, selector):
    # This needs to be updated in the router
    headers = {'Content-type': 'application/json'}
    url = f'{URL}/{db_name}/_find'
    response = requests.post(url, data=json.dumps(
        {'selector': selector}), headers=headers)
    return response.json()

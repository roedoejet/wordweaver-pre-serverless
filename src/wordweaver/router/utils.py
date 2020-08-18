import csv
import os
import re
from tempfile import mkstemp

from docx import Document
from jinja2 import Environment, FileSystemLoader
from latex import build_pdf
from pydantic import BaseModel
from pylatexenc.latexencode import utf8tolatex
from wordweaver.models import Response


class FileSettings(BaseModel):
    heading: str = "Conjugations"
    headers: bool = True


class File:
    """[summary]
    """

    def __init__(self, conjugations: Response = None, tiers=None):
        self._conjugations = conjugations
        self._tiers = tiers
        self.format()

    @property
    def conjugations(self):
        """
        """
        return self._conjugations

    @property
    def tiers(self):
        """
        """
        return self._tiers

    @property
    def formatted_data(self):
        """[summary]
        """
        return self._formatted_data

    def sort_morphemes(self, x):
        if "position" in x:
            return x["position"]
        else:
            return 0

    def format(self):
        """[summary]

        Returns:
            [type]: [description]
        """
        self._formatted_data = []
        for conjugation in self.conjugations:
            tiered_conjugation = []
            for tier in self.tiers:
                tier = tier.dict()
                # Filter empty and sort by position
                output = sorted(
                    [
                        x
                        for x in conjugation["output"]
                        if tier["key"] in x and x[tier["key"]]
                    ],
                    key=self.sort_morphemes,
                )
                # Join with separator
                output = tier["separator"].join(map(lambda x: x[tier["key"]], output))
                tiered_conjugation.append(
                    {"name": tier["name"], "options": tier["options"], "output": output}
                )
            self._formatted_data.append(tiered_conjugation)


class DocxFile(File):
    def __init__(self, conjugations, tiers, settings: FileSettings):
        self.document = Document()
        super(DocxFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()

    def write_to_temp(self):
        # Add header
        self.document.add_heading(self.settings["heading"], 0)
        # Add each tier
        for conjugation in self.formatted_data:
            tiers = "\n".join(map(lambda x: x["output"], conjugation))
            self.document.add_paragraph(tiers, style="List Number")
            self.document.add_paragraph()
        fd, path = mkstemp()
        self.document.save(path)
        return path


class LatexFile(File):
    def __init__(self, conjugations, tiers, settings):
        super(LatexFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()
        self.latexJinjaEnv = Environment(
            block_start_string="\jblock{",  # noqa: W605
            block_end_string="}",
            variable_start_string="\jvar{",  # noqa: W605
            variable_end_string="}",
            comment_start_string="\#{",  # noqa: W605
            comment_end_string="}",
            line_statement_prefix="%%",
            line_comment_prefix="%#",
            trim_blocks=True,
            autoescape=False,
            loader=FileSystemLoader(os.path.dirname(__file__)),
        )
        self.template = self.latexJinjaEnv.get_template("template.tex")

    def write_to_temp(self, ftype="tex"):
        formatted_tiers = [
            [conjugation["output"] for conjugation in x] for x in self.formatted_data
        ]
        data = {"title": self.settings["heading"], "conjugations": formatted_tiers}
        tex = utf8tolatex(
            self.sanitize(self.template.render(data=data)), non_ascii_only=True
        )
        fd, path = mkstemp()
        if ftype == "pdf":
            pdf = build_pdf(tex)
            pdf.save_to(path)
        else:
            with open(path, "w") as f:
                f.write(tex)
        return path

    def sanitize(self, data):
        escape_characters = "".join(["%", "&"])
        findall_pattern = re.compile(r"(?<=[^\\])[{}]".format(escape_characters))
        try:
            matches = findall_pattern.findall(data)
            replaced_data = data
            for match in matches:
                find_pattern = r"(?<=[^\\]){}".format(match)
                replace_pattern = r"\\{}".format(match)
                replaced_data = re.sub(find_pattern, replace_pattern, replaced_data)
            return replaced_data
        except AttributeError:
            return data
        return data


class CsvFile(File):
    def __init__(self, conjugations, tiers, settings: FileSettings):
        super(CsvFile, self).__init__(conjugations, tiers)
        self.settings = settings.dict()

    def write_to_temp(self):
        fd, path = mkstemp()
        with open(path, "w") as f:
            writer = csv.writer(f)
            if self.settings["headers"]:
                writer.writerow([x.dict()["name"] for x in self.tiers])
            for conjugation in self.formatted_data:
                writer.writerow([x["output"] for x in conjugation])
        return path
